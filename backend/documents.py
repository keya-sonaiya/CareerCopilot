from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader


SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentExtractionError(ValueError):
    pass


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError("Could not decode text file.")


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip())


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                table_cells.append(" | ".join(values))

    return "\n".join([*paragraphs, *table_cells])


async def extract_resume_text(upload: UploadFile) -> tuple[str, dict[str, str | int]]:
    filename = upload.filename or "resume"
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_RESUME_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_RESUME_EXTENSIONS))
        raise DocumentExtractionError(f"Unsupported resume file type. Upload one of: {supported}.")

    content = await upload.read()
    if not content:
        raise DocumentExtractionError("Uploaded resume file is empty.")

    try:
        if extension == ".pdf":
            text = _extract_pdf_text(content)
        elif extension == ".docx":
            text = _extract_docx_text(content)
        else:
            text = _decode_text(content)
    except Exception as exc:
        if isinstance(exc, DocumentExtractionError):
            raise
        raise DocumentExtractionError(f"Could not extract text from {filename}.") from exc

    text = text.strip()
    if len(text) < 20:
        raise DocumentExtractionError("Could not extract enough resume text from the uploaded file.")

    return text, {
        "filename": filename,
        "extension": extension,
        "bytes": len(content),
        "characters": len(text),
    }
